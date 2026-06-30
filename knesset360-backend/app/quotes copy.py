import math
import os
import re
import requests
import docx
import pandas as pd
import sqlalchemy
from io import BytesIO
from datetime import datetime
from elasticsearch import Elasticsearch
from elasticsearch.helpers import parallel_bulk
from config import config
import functools
import re


# LOCAL ENGINES 
ES_HOST = "http://localhost:9200"

relevant_tags = ['דובר-המשך', 'קריאות', 'יור', 'דובר', 'קריאה', 'דובר_המשך']
re_remove = re.compile(r'<<.+?>>|/(.+?/)|היו"ר|[^\sא-ת]')
re_space = re.compile(r'\s+')
re_tags = re.compile('<<.+?>>')

db_params = config(filename='../../db-files/database.ini')

es = Elasticsearch(ES_HOST)
engine = sqlalchemy.create_engine(
    f"postgresql+psycopg2://{db_params['user']}:{db_params['password']}@{db_params['host']}:5432/{db_params['database']}"
)


def categorize_text(text):
    if len(text.strip()) < 10: 
        return "General"
    
    prompt = f"Classify this Hebrew text into one word only ('Road Safety', 'Health', 'Education', 'Crime' or 'General'): {text}"
    
    try:
        res = requests.post(OLLAMA_HOST, json={
            "model": "qwen2.5:7b",
            "prompt": prompt,
            "stream": False
        })
        
        # 1. Grab the raw text response
        raw_output = res.json().get("response", "General")
        
        # 2. STRIP hidden spaces/newlines and convert to lowercase!
        clean_output = raw_output.strip().lower()
        
        # 3. Match against lowercase keys, map to clean display values
        mapping = {
            "road safety": "Road Safety",
            "health": "Health",
            "education": "Education",
            "crime": "Crime"
        }
        
        # If it matches our mapping, return the proper casing, otherwise "General"
        return mapping.get(clean_output, "General")
        
    except Exception as e:
        print(f"Connection/Parsing error: {e}")
        return "General"


#  WORD FILE EXTRACTION 
def get_doc_from_url(url):
    try:
        content = requests.get(url, timeout=10).content
        if content.startswith(b"PK"):  # Modern .docx
            return docx.Document(BytesIO(content))
        else:  # Legacy .doc (Requires Windows MS Word Application locally)
            import win32com.client as win32
            filepath = os.path.join(os.getcwd(), 'tmp.doc')
            with open(filepath, 'wb') as f:
                f.write(content)
            word = win32.gencache.EnsureDispatch('Word.Application')
            doc = word.Documents.Open(filepath)
            doc.Activate()
            word.ActiveDocument.SaveAs(filepath, FileFormat=12) # wdFormatXMLDocument
            doc.Close(False)
            buffer = BytesIO(open(filepath, "rb").read())
            os.remove(filepath)
            return docx.Document(buffer)
    except Exception as e:
        print(f"Skipping URL due to parsing error: {url} | Error: {e}")
        return None

def clean_speaker_name(text):
    text = re.sub(r"<<.*?>>", "", text)
    text = re.sub(r"\(.*?\)", "", text)
    text = re.sub(r"היו\"ר\s*", "", text)
    return text.replace(":", "").strip()


# Determine if a paragraph represents a speaker line in the protocol
def is_speaker(p):
    if is_speaker_strong(p):
        return True
    return p.text.endswith(':') and p.runs[0].underline

# Strong detection of speaker paragraphs based on Word style tags
def is_speaker_strong(p):
    return p.text and len(p.text.strip()) > 2 and p.style.name in relevant_tags

def process_protocols(df, doc_type):
    dfs = []
    non_existent = 0
    j = 0
    for i, row in df.iterrows():
        j += 1
        print(f'processing row {j} - {round(j / df.shape[0] * 100, 2)}%, S: {non_existent}', end='\r')
        print(row)
        url = row['filepath']
        document = get_doc_from_url(url)

        # Skip processing loops for files that couldn't be loaded/reached
        if not document:
            non_existent += 1
            continue

        running_index = 0
        latest_speaker = ''
        should_record = False
        records = []
        try:
            for p in document.paragraphs:
                if (not should_record and is_speaker_strong(p)) or (should_record and is_speaker(p)):
                    should_record = True
                    latest_speaker = clean_speaker_name(p.text)
                    continue

                if should_record and len(p.text) > 1 and '<<' not in p.text:
                    running_index += 1
                    records.append(
                        {'Index': running_index, 'Speaker': latest_speaker, 'RawText': p.text, 'DocType': doc_type})
            
            file_df = pd.DataFrame.from_records(records)
            file_df['DocumentId'] = row['documentid']
            file_df['StartDate'] = row['startdate']

            if len(file_df):
                dfs.append(file_df)
        except Exception as e:
            print(f'Error: {e}', url)
            import traceback
            traceback.print_exc()

    df = pd.concat(dfs, ignore_index=True)
    return df


def parse_person(p):
    first_name = p.firstname
    last_name = p.lastname
    s = set()
    s.update(first_name.split())
    s.update(last_name.split())
    
    return {
        'first': first_name,
        'first_split': first_name.split(),
        'last': last_name,
        'last_split': last_name.split(),
        'full_name': f'{first_name} {last_name}',
        'parts': s,
    }

@functools.lru_cache(64000)
def name_to_id(name):
    #if name in parsed_name_to_id_cache:
        #return parsed_name_to_id_cache[name]
    
    name_split = name.split()
    
    if len(name_split) < 2 or len(name_split) > 10:
        return None
    
    for id_, parsed in parsed_person.items():
        # exact match (order doesnt matter)?
        if all(map(lambda n: n in parsed['parts'], name_split)):
            return id_
        
    for id_, parsed in parsed_person.items():
        # partial match (last name equals, first name contains at least one piece)
        if parsed['last'] in name_split and any(map(lambda n: n in name_split, parsed['first_split'])):
            return id_
    
    for id_, parsed in parsed_person.items():
        # full name substring
        if parsed['full_name'] in name:
            return id_
    """
    for id_, parsed in parsed_person.items():
        # split name substring
        if parsed['first'] in name and parsed['last'] in name:
            return id_
            """

def clean_speaker(txt):
    if pd.isna(txt):
        return ''
    txt = str(txt)

    txt = txt.replace('יי', 'י')
    txt = txt.replace('וו', 'ו')
    txt = txt.replace('אא', 'א')

    txt = re_remove.sub('', txt)
    txt = re_space.sub(' ', txt)
    txt = re_tags.sub('', txt)

    return txt.strip()


#  EXECUTION MATRIX 
if __name__ == "__main__":
    # Initialize Hebrew search capabilities inside Elasticsearch
    if not es.indices.exists(index="quotes"): 
        es.indices.create(index="quotes", body={
            "mappings": {
                "properties": {
                    "Speaker": {"type": "keyword"},
                    "RawText": {"type": "text" },
                    "DocType": {"type": "keyword"},
                    "DocumentId": {"type": "integer"},
                    "StartDate": {"type": "date"}
                }
            }
        })

    print("Pulling documents from PostgreSQL...")
    
    # Load person table
    mks = pd.read_sql(
        """
        SELECT p.id, p.firstname, p.lastname FROM KNS_Person p
        JOIN KNS_PersonToPosition kptp on kptp.personid = p.id
        WHERE kptp.StartDate > '2000-01-01'
        GROUP BY p.id, p.firstname, p.lastname
        ORDER BY MAX(kptp.StartDate) DESC
        """,
        engine)

    mks['CleanedFirstName'] = mks['firstname'].apply(clean_speaker)
    mks['CleanedLastName'] =  mks['lastname'].apply(clean_speaker)
    parsed_person = {}
    for _, p in mks.iterrows():
        parsed_person[p.id] = parse_person(p)


    sql_urls_comittees = """
        SELECT s.StartDate, d.id as documentid, d.FilePath
        FROM KNS_CommitteeSession s
        JOIN KNS_DocumentCommitteeSession d on d.CommitteeSessionID = s.id
        WHERE d.ApplicationID = 1
        AND d.GroupTypeID = 23
        AND s.StartDate >= Timestamp '15/06/2026'
        ORDER BY s.StartDate ASC
    """

    sql_urls_plenums = """
        SELECT s.StartDate, d.id as documentid, d.FilePath
        FROM KNS_PlenumSession s
        JOIN KNS_DocumentPlenumSession d on d.PlenumSessionID = s.id
        WHERE d.ApplicationID = 1
        AND d.GroupTypeID = 28
        AND s.StartDate >= Timestamp '15/06/2026'
        ORDER BY s.StartDate ASC
    """
    finished_urls = 0

    for chunk_df in pd.read_sql(sql_urls_comittees, engine, chunksize=100):

        quotes = process_protocols(chunk_df, doc_type="Committee")
        if quotes.empty:
            continue

        speakers = quotes["Speaker"].apply(clean_speaker).unique()
        # Map extracted speaker name to person ID
        name_mapping = {}
        total = len(speakers)

        for i, name in enumerate(speakers):
            try:
                name_mapping[name] = name_to_id(name)
            except Exception:
                name_mapping[name] = None

            print(f'{i}, {round((i+1) / total * 100, 2)}%      ', end='\r')

        # Apply name mapping
        quotes['PersonID'] = (quotes['Speaker'].apply(clean_speaker).map(name_mapping))
        quotes.replace({math.nan: None}, inplace=True)
        
        def gen():
            j = 0
            for i, row in quotes.iterrows():
                j = j + 1
                row = dict(row)
                row["$Timestamp"] = datetime.now()
                row["_index"] = "quotes"
                row["_id"] = f'{row["DocType"]}:{row["DocumentId"]}:{row["Index"]}'
                yield row
                if j % 1000 == 0:
                    print(f'Pushed {j} lines          ', end='\r')
            
        for success, info in parallel_bulk(es, gen()):
            if not success:
                print('failed', info)

        finished_urls += 100
        print(f'Finished {finished_urls} urls')

    # es.indices.delete(index="quotes")